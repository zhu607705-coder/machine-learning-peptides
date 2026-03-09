from __future__ import annotations

import html
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as PdfImage,
    PageBreak,
    PageTemplate,
    Paragraph,
    Preformatted,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "多肽合成机器学习学习报告.md"
OUTPUT_DOCX = ROOT / "报告" / "多肽合成机器学习学习报告.docx"
OUTPUT_PDF = ROOT / "报告" / "多肽合成机器学习学习报告.pdf"
TMP_RENDER_DIR = ROOT / "tmp" / "pdfs" / "learning-report-pages"
AUTHOR_NAME = "朱航成"
AUTHOR_ID = "3250100360"
AUTHOR_COLLEGE = "云峰学园"
AUTHOR_MAJOR = "化学工程与工艺"


@dataclass
class Block:
    kind: str
    node: Tag | None = None
    text: str | None = None


def ensure_dirs() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    TMP_RENDER_DIR.mkdir(parents=True, exist_ok=True)


def preprocess_markdown(text: str) -> str:
    lines = text.splitlines()
    filtered: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("**副标题：**"):
            continue
        if stripped.startswith("**作者：**"):
            continue
        if stripped.startswith("**日期：**"):
            continue
        if stripped.startswith("**关键词：**"):
            continue
        filtered.append(line)
    text = "\n".join(filtered)
    text = text.replace("### 一、学术论文", "## 参考文献\n\n### 一、学术论文")
    text = re.sub(r"<sup>\[([^\]]+)\]</sup>", r"[\1]", text)
    return text


def markdown_to_blocks(text: str) -> list[Block]:
    html_text = markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "sane_lists"],
        output_format="html5",
    )
    soup = BeautifulSoup(html_text, "html.parser")
    blocks: list[Block] = []
    for node in soup.contents:
        if isinstance(node, NavigableString):
            continue
        if not isinstance(node, Tag):
            continue
        if node.name == "hr":
            blocks.append(Block(kind="pagebreak"))
            continue
        if node.name == "pre":
            blocks.append(Block(kind="pre", text=node.get_text("\n", strip=False)))
            continue
        blocks.append(Block(kind=node.name, node=node))
    return blocks


def extract_toc_entries(blocks: list[Block]) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    for block in blocks:
        if block.kind not in {"h1", "h2"}:
            continue
        text = block.node.get_text(" ", strip=True)
        if text in {"多肽合成机器学习学习报告", "参考文献"}:
            continue
        if block.kind == "h1":
            entries.append((1, text))
        elif block.kind == "h2":
            entries.append((2, text))
    return entries


class ZjuReportDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, **kwargs) -> None:
        super().__init__(filename, **kwargs)
        frame = Frame(
            self.leftMargin,
            self.bottomMargin,
            self.width,
            self.height,
            id="normal",
        )
        self.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=self._draw_running)])

    def _draw_running(self, canvas, doc) -> None:
        if canvas.getPageNumber() > 1:
            canvas.setFont("STSong-Light", 9)
            canvas.setFillColor(colors.HexColor("#666666"))
            canvas.drawCentredString(A4[0] / 2, A4[1] - 1.2 * cm, "浙江大学学习报告")
            canvas.line(3.0 * cm, A4[1] - 1.45 * cm, A4[0] - 2.5 * cm, A4[1] - 1.45 * cm)
            canvas.drawCentredString(A4[0] / 2, 1.2 * cm, f"{canvas.getPageNumber()}")

    def afterFlowable(self, flowable) -> None:
        return


def set_run_fonts(run, east_asia: str = "Songti SC", western: str = "Times New Roman") -> None:
    run.font.name = western
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_text(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    set_run_fonts(run)
    run.bold = bold
    run.italic = italic


def apply_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")

    for level, size in [(1, 16), (2, 15), (3, 14), (4, 12)]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        style.font.size = Pt(size)
        style.font.bold = True

    if "CodeBlock" not in document.styles:
        style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        style.font.size = Pt(9.5)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char_begin)
    run._element.append(instr)
    run._element.append(fld_char_end)
    set_run_fonts(run)


def configure_docx_running_elements(section, header_text: str) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = ""
    run = header.add_run(header_text)
    set_run_fonts(run)
    run.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    add_page_number_field(footer)
    if footer.runs:
        footer.runs[0].font.size = Pt(10.5)


def docx_add_title_page(document: Document) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True
    p0 = document.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(18)
    run = p0.add_run("浙江大学学习报告")
    set_run_fonts(run)
    run.font.size = Pt(18)
    run.bold = True

    p1 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(80)
    run = p1.add_run("多肽合成机器学习项目")
    set_run_fonts(run)
    run.font.size = Pt(16)
    run.bold = True

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(16)
    run = title.add_run("多肽合成机器学习学习报告")
    set_run_fonts(run)
    run.font.size = Pt(22)
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(60)
    run = subtitle.add_run("从泛化化工过程回归到基于 Fmoc SPPS SOP 的序列合成难度与风险分析")
    set_run_fonts(run)
    run.font.size = Pt(13)
    run.italic = True

    meta_lines = [
        f"姓名：{AUTHOR_NAME}",
        f"学号：{AUTHOR_ID}",
        f"学院：{AUTHOR_COLLEGE}",
        f"专业：{AUTHOR_MAJOR}",
        "日期：2026-03-03",
        "关键词：多肽合成；机器学习；Fmoc SPPS；分层模型；过程优化",
    ]
    for line in meta_lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(line)
        set_run_fonts(run)
        run.font.size = Pt(12)


def docx_add_toc_page(document: Document, entries: list[tuple[int, str]]) -> None:
    p = document.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(p, "目 录", bold=True)
    for level, text in entries:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if level == 2:
            p.paragraph_format.left_indent = Cm(0.7)
        dots = "." * max(10, 34 - len(text))
        set_paragraph_text(p, f"{text} {dots}")
    document.add_page_break()


def add_docx_table(document: Document, node: Tag) -> None:
    rows = []
    header = [cell.get_text(" ", strip=True) for cell in node.find("tr").find_all(["th", "td"])]
    rows.append(header)
    for tr in node.find_all("tr")[1:]:
        rows.append([cell.get_text(" ", strip=True) for cell in tr.find_all(["th", "td"])])
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_fonts(run)
            run.font.size = Pt(10.5)
            if r_idx == 0:
                run.bold = True
    document.add_paragraph()


def add_docx_list(document: Document, node: Tag, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in node.find_all("li", recursive=False):
        p = document.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        set_paragraph_text(p, li.get_text(" ", strip=True))


def add_docx_image(document: Document, node: Tag) -> None:
    src = node.get("src")
    if not src:
        return
    image_path = (SOURCE_MD.parent / src).resolve()
    if not image_path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(14.5))


def export_docx(blocks: list[Block]) -> None:
    document = Document()
    apply_docx_styles(document)
    docx_add_title_page(document)
    body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    configure_docx_running_elements(body_section, "浙江大学学习报告")

    first_h1_seen = False
    toc_inserted = False
    toc_entries = extract_toc_entries(blocks)
    abstract_seen = False
    for block in blocks:
        if block.kind == "h1":
            text = block.node.get_text(" ", strip=True)
            if text == "多肽合成机器学习学习报告":
                continue
            if not first_h1_seen and text == "参考文献":
                first_h1_seen = True
            p = document.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_text(p, text, bold=True)
        elif block.kind == "h2":
            text = block.node.get_text(" ", strip=True)
            if text == "摘要":
                abstract_seen = True
            if abstract_seen and not toc_inserted:
                if text != "摘要":
                    docx_add_toc_page(document, toc_entries)
                    toc_inserted = True
            p = document.add_paragraph(style="Heading 2")
            set_paragraph_text(p, text, bold=True)
        elif block.kind == "h3":
            p = document.add_paragraph(style="Heading 3")
            set_paragraph_text(p, block.node.get_text(" ", strip=True), bold=True)
        elif block.kind == "h4":
            p = document.add_paragraph(style="Heading 4")
            set_paragraph_text(p, block.node.get_text(" ", strip=True), bold=True)
        elif block.kind == "p":
            if block.node.find("img"):
                add_docx_image(document, block.node.find("img"))
                continue
            p = document.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(4)
            set_paragraph_text(p, block.node.get_text(" ", strip=True))
        elif block.kind == "ul":
            add_docx_list(document, block.node, ordered=False)
        elif block.kind == "ol":
            add_docx_list(document, block.node, ordered=True)
        elif block.kind == "table":
            add_docx_table(document, block.node)
        elif block.kind == "pre":
            p = document.add_paragraph(style="CodeBlock")
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(4)
            set_paragraph_text(p, block.text.rstrip("\n"))
        elif block.kind == "pagebreak":
            document.add_page_break()

    document.save(OUTPUT_DOCX)


def build_pdf_styles():
    registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CJKBody",
            fontName="STSong-Light",
            fontSize=11,
            leading=18,
            alignment=TA_JUSTIFY,
            spaceAfter=7,
            firstLineIndent=22,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CJKTitle",
            fontName="STSong-Light",
            fontSize=22,
            leading=30,
            alignment=TA_CENTER,
            spaceAfter=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CJKSubtitle",
            fontName="STSong-Light",
            fontSize=12,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CJKMeta",
            fontName="STSong-Light",
            fontSize=11,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=6,
        )
    )
    for name, size in [("CJKH1", 16), ("CJKH2", 14), ("CJKH3", 12.5), ("CJKH4", 11.5)]:
        styles.add(
            ParagraphStyle(
                name=name,
                fontName="STSong-Light",
                fontSize=size,
                leading=size + 6,
                alignment=TA_LEFT,
                spaceBefore=10,
                spaceAfter=8,
            )
        )
    styles.add(
        ParagraphStyle(
            name="CJKCaption",
            fontName="STSong-Light",
            fontSize=10,
            leading=14,
            alignment=TA_CENTER,
            spaceBefore=2,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CJKCode",
            fontName="STSong-Light",
            fontSize=9.2,
            leading=14,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CJKList",
            fontName="STSong-Light",
            fontSize=11,
            leading=17,
            leftIndent=16,
            firstLineIndent=0,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCHeading",
            fontName="STSong-Light",
            fontSize=16,
            leading=22,
            alignment=TA_LEFT,
            spaceBefore=0,
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCLevel1",
            fontName="STSong-Light",
            fontSize=11.5,
            leading=18,
            leftIndent=0,
            firstLineIndent=0,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCLevel2",
            fontName="STSong-Light",
            fontSize=10.8,
            leading=17,
            leftIndent=16,
            firstLineIndent=0,
            spaceAfter=2,
        )
    )
    return styles


def pdf_escape(text: str) -> str:
    return html.escape(text).replace("\n", "<br/>")


def add_pdf_title_page(story: list, styles) -> None:
    story.append(Spacer(1, 2.2 * cm))
    story.append(Paragraph("浙江大学学习报告", styles["CJKTitle"]))
    story.append(Paragraph("多肽合成机器学习项目", styles["CJKSubtitle"]))
    story.append(Spacer(1, 1.8 * cm))
    story.append(Paragraph("多肽合成机器学习学习报告", styles["CJKTitle"]))
    story.append(Paragraph("从泛化化工过程回归到基于 Fmoc SPPS SOP 的序列合成难度与风险分析", styles["CJKSubtitle"]))
    story.append(Spacer(1, 0.8 * cm))
    for line in [
        f"姓名：{AUTHOR_NAME}",
        f"学号：{AUTHOR_ID}",
        f"学院：{AUTHOR_COLLEGE}",
        f"专业：{AUTHOR_MAJOR}",
        "日期：2026-03-03",
        "关键词：多肽合成；机器学习；Fmoc SPPS；分层模型；过程优化",
    ]:
        story.append(Paragraph(line, styles["CJKMeta"]))
    story.append(PageBreak())


def normalize_for_match(text: str) -> str:
    return re.sub(r"\s+", "", text)


def build_static_toc_page(styles, entries: list[tuple[int, str]], page_map: dict[str, int]) -> list:
    story: list = [Paragraph("目 录", styles["TOCHeading"])]
    for level, text in entries:
        page_no = page_map.get(text, "")
        dots = "." * max(8, 48 - len(text) - len(str(page_no)))
        line = f"{text} {dots} {page_no}"
        style = styles["TOCLevel1"] if level == 1 else styles["TOCLevel2"]
        story.append(Paragraph(pdf_escape(line), style))
    story.append(PageBreak())
    return story


def find_heading_pages(pdf_path: Path, entries: list[tuple[int, str]]) -> dict[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(str(pdf_path))
    page_map: dict[str, int] = {}
    normalized_targets = {text: normalize_for_match(text) for _, text in entries}
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = normalize_for_match(page.extract_text() or "")
        for text, target in normalized_targets.items():
            if text in page_map:
                continue
            if target and target in page_text:
                page_map[text] = page_index
    if "摘要" in page_map:
        abstract_page = page_map["摘要"]
        for text in list(page_map):
            if text != "摘要" and page_map[text] >= abstract_page:
                page_map[text] += 1
    return page_map


def html_children(node: Tag) -> Iterable[Tag | NavigableString]:
    for child in node.children:
        yield child


def list_texts(node: Tag) -> list[str]:
    return [li.get_text(" ", strip=True) for li in node.find_all("li", recursive=False)]


def add_pdf_table(story: list, node: Tag, styles) -> None:
    rows = []
    first = node.find("tr")
    rows.append([Paragraph(pdf_escape(cell.get_text(" ", strip=True)), styles["CJKCaption"]) for cell in first.find_all(["th", "td"])])
    for tr in node.find_all("tr")[1:]:
        rows.append([Paragraph(pdf_escape(cell.get_text(" ", strip=True)), styles["CJKBody"]) for cell in tr.find_all(["th", "td"])])
    col_count = len(rows[0])
    table = Table(rows, repeatRows=1, hAlign="CENTER", colWidths=[16.5 * cm / col_count] * col_count)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef3f8")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#7a8699")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.22 * cm))


def build_pdf_story(blocks: list[Block], styles, toc_page_map: dict[str, int] | None = None) -> list:
    story = []
    add_pdf_title_page(story, styles)
    toc_inserted = False
    abstract_seen = False

    for block in blocks:
        if block.kind == "h1":
            text = block.node.get_text(" ", strip=True)
            if text == "多肽合成机器学习学习报告":
                continue
            story.append(Paragraph(pdf_escape(text), styles["CJKH1"]))
        elif block.kind == "h2":
            text = block.node.get_text(" ", strip=True)
            if text == "摘要":
                abstract_seen = True
            elif abstract_seen and not toc_inserted and toc_page_map is not None:
                story.extend(build_static_toc_page(styles, extract_toc_entries(blocks), toc_page_map))
                toc_inserted = True
            story.append(Paragraph(pdf_escape(text), styles["CJKH2"]))
        elif block.kind == "h3":
            story.append(Paragraph(pdf_escape(block.node.get_text(" ", strip=True)), styles["CJKH3"]))
        elif block.kind == "h4":
            story.append(Paragraph(pdf_escape(block.node.get_text(" ", strip=True)), styles["CJKH4"]))
        elif block.kind == "p":
            img = block.node.find("img")
            if img:
                src = img.get("src")
                image_path = (SOURCE_MD.parent / src).resolve()
                if image_path.exists():
                    story.append(PdfImage(str(image_path), width=15 * cm, height=8.5 * cm, hAlign="CENTER"))
                    story.append(Spacer(1, 0.15 * cm))
                continue
            story.append(Paragraph(pdf_escape(block.node.get_text(" ", strip=True)), styles["CJKBody"]))
        elif block.kind == "ul":
            for text in list_texts(block.node):
                story.append(Paragraph(pdf_escape(f"- {text}"), styles["CJKList"]))
            story.append(Spacer(1, 0.12 * cm))
        elif block.kind == "ol":
            for idx, text in enumerate(list_texts(block.node), start=1):
                story.append(Paragraph(pdf_escape(f"{idx}. {text}"), styles["CJKList"]))
            story.append(Spacer(1, 0.12 * cm))
        elif block.kind == "table":
            add_pdf_table(story, block.node, styles)
        elif block.kind == "pre":
            story.append(Preformatted(block.text.rstrip("\n"), styles["CJKCode"]))
        elif block.kind == "pagebreak":
            story.append(PageBreak())
    return story


def build_pdf_file(output_path: Path, story: list) -> None:
    doc = ZjuReportDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=3.0 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.54 * cm,
        bottomMargin=2.54 * cm,
        title="多肽合成机器学习学习报告",
        author=AUTHOR_NAME,
    )
    doc.build(story)


def export_pdf(blocks: list[Block]) -> None:
    styles = build_pdf_styles()
    temp_pdf = TMP_RENDER_DIR / "report-pass1.pdf"
    build_pdf_file(temp_pdf, build_pdf_story(blocks, styles, toc_page_map=None))
    page_map = find_heading_pages(temp_pdf, extract_toc_entries(blocks))
    build_pdf_file(OUTPUT_PDF, build_pdf_story(blocks, styles, toc_page_map=page_map))
    if temp_pdf.exists():
        temp_pdf.unlink()


def main() -> None:
    ensure_dirs()
    text = preprocess_markdown(SOURCE_MD.read_text(encoding="utf-8"))
    blocks = markdown_to_blocks(text)
    export_docx(blocks)
    export_pdf(blocks)


if __name__ == "__main__":
    main()
